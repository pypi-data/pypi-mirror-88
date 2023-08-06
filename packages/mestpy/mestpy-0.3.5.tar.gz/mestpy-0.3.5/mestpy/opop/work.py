def docxToPdf(directory):
  from docx2pdf import convert
  import os
  files = os.listdir(directory)
  for ids,f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    name = '{} {} {}'.format(ids+1,f_name, f_ext)
    print(name)
  convert(directory+'/')
  print("pdf файлы сохранены в папке c docx")

def check_tables(directory, numTable):
  from docx import Document
  import os
  import copy
  files = os.listdir(directory)
  listSortedNatural = files
  document_out = Document() # новый документ с таблицами 0 всех файлов
  text1 = 'таблицы {} из всех РПД'.format(numTable)
  document_out.add_paragraph(text1)
  for ids, rpd in enumerate(listSortedNatural): 
    f_name, f_ext = os.path.splitext(rpd)
    if f_ext != ".docx":
      print("Не docx file: {}".format(rpd))
      continue
    text2 = 'Файл '+str(ids+1)
    document_out.add_paragraph(text2)
    document_out.add_paragraph(rpd)  #print(rpd)
    doc_document = os.path.join(directory, rpd)
    document = Document(doc_document)
    for k, table in enumerate(document.tables):
      if k == numTable-1:
        template0 = document.tables[numTable]
        tbl0 = template0._tbl
        new_tbl0 = copy.deepcopy(tbl0)
        paragraph1 = document_out.add_paragraph()
        paragraph1._p.addnext(new_tbl0)
    document_out.add_page_break()
  document_out.save('all_tables.docx')
  print("создал файл all_tables.docx")

def fromRupFullNames(fileRUP):
  import pandas as pd
  df_rup = pd.read_excel(fileRUP, sheet_name = 'Компетенции(2)', index_col = None, header = None )
  df_rup = df_rup.dropna(subset = [4])
  df_rup.drop(df_rup.head(2).index,inplace = True) # drop last n rows 
  df_rup.drop(df_rup.tail(12).index,inplace = True) # drop last n rows 
  df_rup.dropna(axis = 'columns',how = 'all', inplace = True)
  df_rup.drop(df_rup.columns[[0,1, 3,5]], axis = 1, inplace = True)
  df_rup.reset_index(drop = True, inplace = True)
  df_rup = df_rup.rename(columns = {2: "Index", 4: "Name"})
  writer = pd.ExcelWriter("rupFullNames" + ".xlsx", engine = 'xlsxwriter')
  df_rup.to_excel(writer, sheet_name = 'FullName', startrow = 0, startcol = 0, index = False)
  writer.save()
  print("сохранил файл rupFullNames.xlsx на корневой директории py")


def unionLeftRight(fileForUnion):
  import pandas as pd
  df_rup = pd.read_excel(fileForUnion, usecols = "A,B" )
  dfIn = pd.read_excel(fileForUnion)
  dfu = pd.DataFrame
  dfu = dfIn['Left']+" "+dfIn['Right']
  writer = pd.ExcelWriter("unionLeftRightOut" + ".xlsx", engine = 'xlsxwriter')
  dfu.to_excel(writer)
  writer.save()
  print("Создал Файл unionLeftRightOut.xlsx в корневом каталоге Питон")

def fromRupFullNames(fileRUP):
  import pandas as pd
  df_rup = pd.read_excel(fileRUP, sheet_name = 'Компетенции(2)', index_col = None, header = None )
  df_rup = df_rup.dropna(subset = [4])
  df_rup.drop(df_rup.head(2).index,inplace = True) # drop last n rows 
  df_rup.drop(df_rup.tail(12).index,inplace = True) # drop last n rows 
  df_rup.dropna(axis = 'columns',how = 'all', inplace = True)
  df_rup.drop(df_rup.columns[[0,1, 3,5]], axis = 1, inplace = True)
  df_rup.reset_index(drop = True, inplace = True)
  # df_rup.columns = ['code','id','text']
  df_rup = df_rup.rename(columns = {2: "Index", 4: "Name"})
  writer = pd.ExcelWriter("rupFullNames" + ".xlsx", engine = 'xlsxwriter')
  df_rup.to_excel(writer, sheet_name = 'FullName', startrow = 0, startcol = 0, index = False)
  writer.save()
  print("сохранил файл rupFullNames.xlsx на корневой директории py")

def pagesFromPdfToExcel(directoryPdfs):
  from PyPDF2 import PdfFileReader, PdfFileWriter
  import os
  import pandas as pd
  files = os.listdir(directoryPdfs)
  listSortedNatural = files
  lNames = [] 
  lPages = []
  for idx, f in enumerate(listSortedNatural):
    pdf = PdfFileReader(directoryPdfs+"/"+f)
    page = pdf.numPages
    lPages.append(page)
    f_name, f_ext = os.path.splitext(f)
    lNames.append(f_name)
  dfNamesFiles = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel['Кол-во листов'] = lPages
  writer = pd.ExcelWriter("pagesOfPdfFiles" + ".xlsx", engine = 'xlsxwriter')
  dfToExcel.to_excel(writer, startrow = 0, startcol = 0, index = False)
  writer.save()
  print("Сохранил файл pagesOfPdfFiles.xlsx на корневом каталоге py")

import os
def renameAllFiles(directory,inStrins,outString,numStart,numEnd):
    files = os.listdir(directory)
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        print("Old Name {}".format(f_name))
        f_name=f_name.replace(inStrins,outString).strip()
        f_start=f_name
        if numStart>0:
            f_start=f_name[:numStart]
        f_end=""
        if numEnd>0:
            f_end=f_name[len(f_name)-numEnd:]
        f_name=f_start+f_end
        print("New Name {}".format(f_name))
        new_name='{}{}'.format(f_name, f_ext)
        os.rename(directory+'/'+f, directory+'/'+new_name)


def indexForFiles(directory,constStr,startNum):
  import os
  files = os.listdir(directory)
  for ids,f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    print("Old Name {}".format(f_name))
    new_name = '{}{}'.format(f_name, f_ext)
    new_ids = str(ids+startNum)
    new_name = '{}{} {}{}'.format(constStr,new_ids,f_name, f_ext)
    print("New Name {}".format(new_name))
    os.rename(directory+'/'+f, directory+'/'+new_name)
    
def unionLeftRight(fileForUnion):
  import pandas as pd
  df_rup = pd.read_excel(fileForUnion, usecols = "A,B" )
  dfIn = pd.read_excel(fileForUnion)
  dfu = pd.DataFrame
  dfu = dfIn['Left']+" "+dfIn['Right']
  writer = pd.ExcelWriter("unionLeftRightOut" + ".xlsx", engine = 'xlsxwriter')
  dfu.to_excel(writer)
  writer.save()
  print("Создал Файл unionLeftRightOut.xlsx в корневом каталоге Питон")
