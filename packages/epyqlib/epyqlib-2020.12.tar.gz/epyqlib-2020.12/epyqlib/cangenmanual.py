import collections
import itertools
import logging
import os
import signal
import time

import attr
import canmatrix.formats
import click
import docx
import docx.enum.section
import docx.enum.table
import docx.enum.text
import docx.oxml
import docx.oxml.ns
import docx.shared
import docx.table
import lxml.etree

import epyqlib.utils.general

# See file COPYING in this source tree
__copyright__ = "Copyright 2017, EPC Power Corp."
__license__ = "GPLv2+"


logger = logging.getLogger(__name__)


# https://github.com/python-openxml/python-docx/issues/322#issuecomment-265018856
def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = docx.oxml.OxmlElement("w:tblHeader")
    tblHeader.set(docx.oxml.ns.qn("w:val"), "true")
    trPr.append(tblHeader)
    return row


# https://github.com/python-openxml/python-docx/issues/245#issuecomment-208476933
def prevent_row_breaks(table):
    tags = table._element.xpath("//w:tr")
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = docx.oxml.OxmlElement("w:cantSplit")  # Create arbitrary tag
        tag.append(child)  # Append in the new tag


def w(s):
    return "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + s


def shd_element(fill):
    e = lxml.etree.Element(w("shd"))
    # <w:shd w:val="clear" w:color="auto" w:fill="D9D9D9" w:themeFill="background1" w:themeFillShade="D9"/>

    e.attrib[w("fill")] = fill

    return e


def shade(cell, fill):
    cell._tc.tcPr.append(shd_element(fill=fill))


@attr.s
class Table:
    title = attr.ib()
    headings = attr.ib()
    widths = attr.ib()
    total_width = attr.ib(default=10)
    comment = attr.ib(default="")
    rows = attr.ib(default=attr.Factory(list))

    def fill_docx(
        self,
        table,
        title_style=None,
        heading_style=None,
        contents_style=None,
    ):
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        # table.autofit = True
        # table.style = 'CAN Table Base'

        title_present = self.title is not None and len(self.title) > 0

        if title_present:
            row = table.add_row()
            row.cells[0].text = self.title
            title_paragraph = row.cells[0].paragraphs[0]
            title_paragraph.paragraph_format.keep_with_next = True
            if title_style is not None:
                title_paragraph.style = title_style
            shade(row.cells[0], fill="000000")

        if len(self.comment) > 0:
            row = table.add_row()
            row.cells[0].merge(row.cells[-1])
            row.cells[0].text = self.comment
            row.cells[0].paragraphs[0].paragraph_format.keep_with_next = True

        row = table.add_row()
        for cell, heading in zip(row.cells, self.headings):
            cell.text = heading
            if heading_style is not None:
                cell.paragraphs[0].style = heading_style
            cell.paragraphs[0].paragraph_format.keep_with_next = True

            if not title_present:
                shade(cell, fill="000000")

        shadings = (
            {"fill": "D9D9D9"},
            None,
        )

        shadings_iterator = itertools.cycle(shadings)
        if not title_present:
            next(shadings_iterator)

        zipped = zip(self.rows, shadings_iterator)
        total = len(self.rows)
        checkpoint = time.monotonic()
        n = 10
        for i, (r, shading) in enumerate(zipped):
            if i > 0 and i % n == 0:
                now = time.monotonic()
                delta = now - checkpoint
                checkpoint = now
                print("{} / {} ({:.3})".format(i, total, delta / n))

            row = table.add_row()
            for cell, text in zip(row.cells, r):
                cell.text = str(text)
                if shading is not None:
                    shade(cell, **shading)
                if contents_style is not None:
                    cell.paragraphs[0].style = contents_style
                cell.paragraphs[0].paragraph_format.keep_with_next = True

        remaining_width = self.total_width - sum(
            w for w in self.widths if w is not None
        )
        each_width = remaining_width / sum(1 for w in self.widths if w is None)
        widths = [each_width if w is None else w for w in self.widths]
        widths = [w if w is None else docx.shared.Inches(w) for w in widths]
        for column, width in zip(table.columns, widths):
            column.width = width

        total = len(self.rows)
        checkpoint = time.monotonic()
        n = 10
        for i, row in enumerate(table.rows):
            if i > 0 and i % n == 0:
                now = time.monotonic()
                delta = now - checkpoint
                checkpoint = now
                print("{} / {} ({:.3})".format(i, total, delta / n))

            for cell, width in zip(row.cells, widths):
                cell.width = width

        if title_present:
            table.rows[0].cells[0].merge(table.rows[0].cells[-1])


def id_string(message_id):
    return "0x{:08X}".format(message_id)


def tabulate_signals(signals):
    rows = []

    for signal in signals:
        # if len(signal.unit) == 0 and signal.factor != 1:
        # table.append(('***',) * 5)
        startbit = signal.start_bit
        rows.append(
            (
                "",
                "",
                signal.name,
                "{}/{}".format(startbit % 8, startbit // 8),
                signal.size,
                signal.factor,
                signal.unit,
                "{}: {}".format(signal.enumeration, signal.values)
                if signal.enumeration is not None
                else "",
            )
        )
        if len(signal.comment) > 0:
            rows.append((*(("",) * 7), signal.comment))

    return rows


def doc_signals(signals):
    rows = []

    for signal in signals:
        startbit = signal.start_bit
        enumeration = signal.enumeration if signal.enumeration is not None else ""
        rows.append(
            (
                signal.name,
                "{}/{}".format(startbit % 8, startbit // 8),
                signal.size,
                signal.factor,
                signal.unit,
                enumeration,
                signal.comment,
            )
        )

    return rows


@click.command()
@click.option("--can", "-c", type=click.File("rb"), required=True)
@click.option("--template", "-t", type=click.File("rb"), required=True)
@click.option("--output", "-o", type=click.File("wb"), required=True)
@click.option("--verbose", "-v", count=True)
def main(can, template, output, verbose):
    if verbose >= 1:
        logger.setLevel(logging.DEBUG)

    (matrix,) = canmatrix.formats.load(
        can, os.path.splitext(can.name)[1].lstrip(".")
    ).values()

    mux_table_header = (
        "Mux Name",
        "Mux ID",
        "Name",
        "Start",
        "Length",
        "Factor",
        "Units",
        "Enumeration",
        "Comment",
    )
    mux_table = epyqlib.utils.general.TextTable()
    mux_table.append(mux_table_header)
    mux_table_header = mux_table_header[2:]

    frame_table_header = (
        "Frame",
        "ID",
        "Name",
        "Start",
        "Length",
        "Factor",
        "Units",
        "Enumeration",
        "Comment",
    )

    frame_table = epyqlib.utils.general.TextTable()
    frame_table.append(frame_table_header)
    frame_table_header = frame_table_header[2:]

    widths = [0.625] * len(frame_table_header)
    widths[0] = 2
    widths[-2] = 1.5
    widths[-1] = None

    enum_table_header = ("Value", "Name")

    frame_tables = []
    multiplex_tables = []
    enumeration_tables = []

    for frame in sorted(matrix.frames, key=lambda f: f.name):
        message_id = frame.arbitration_id.id
        frame_table.append(frame.name, id_string(message_id))

        a_ft = Table(
            title="{} ({})".format(frame.name, id_string(message_id)),
            comment=frame.comment,
            headings=frame_table_header,
            widths=widths,
        )
        frame_tables.append(a_ft)

        mux_table.append("{} ({})".format(frame.name, id_string(message_id)))

        multiplex_signal = frame.signals[0]
        if multiplex_signal.multiplex is None:
            multiplex_signal = None

        if multiplex_signal is None:
            signals = sorted(frame.signals, key=lambda s: s.name)
            table = tabulate_signals(signals)
            frame_table.extend(table)

            a_ft.rows.extend(doc_signals(signals))
        else:
            signals = (multiplex_signal,)
            table = tabulate_signals(signals)
            frame_table.extend(table)
            a_ft.rows.extend(doc_signals(signals))
            # multiplex_signal.values = {
            #     int(k): v for k, v in multiplex_signal.values.items()
            # }

            for value, name in sorted(multiplex_signal.values.items()):
                if value == 0:
                    continue

                a_mt = Table(
                    title="{} ({}) - {} ({})".format(
                        frame.name,
                        id_string(message_id),
                        name,
                        value,
                    ),
                    # comment=frame.comment,
                    headings=mux_table_header,
                    widths=widths,
                )
                multiplex_tables.append(a_mt)

                mux_table.append(
                    "",
                    value,
                    name,
                )

                signals = tuple(
                    s
                    for s in frame.signals
                    if s.multiplex == value and s.name != "ReadParam_command"
                )

                mux_table.extend(
                    tabulate_signals(sorted(signals, key=lambda s: s.name))
                )

                a_mt.rows.extend(doc_signals(signals))

    print("\n\n - - - - - - - Multiplexes\n")
    print(mux_table)

    print("\n\n - - - - - - - Frames\n")
    print(frame_table)

    enumeration_table = epyqlib.utils.general.TextTable()
    enumeration_table.append("Name", "", "Value")
    widths = (1.5, None)
    for name, values in sorted(matrix.value_tables.items()):
        a_et = Table(
            title=name,
            headings=enum_table_header,
            widths=widths,
            total_width=5,
        )
        enumeration_tables.append(a_et)
        a_et.rows.extend(sorted(values.items()))

        enumeration_table.append(name)
        for i, s in sorted(values.items()):
            enumeration_table.append("", s, i)

    print("\n\n - - - - - - - Enumerations\n")
    print(enumeration_table)

    doc = docx.Document(template)

    table_sets = {
        "frames": frame_tables,
        "multiplexers": multiplex_tables,
        "enumerations": enumeration_tables,
    }
    for tag, tables in table_sets.items():
        full_tag = "<gen_{}>".format(tag)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == full_tag:
                break
            elif full_tag in paragraph.text:
                f = "Tag {} found, expected as only text in paragraph: {}"
                raise Exception(f.format(full_tag, repr(paragraph.text)))
        else:
            raise Exception("Tag not found: {}".format(full_tag))

        for table in tables:
            doc_table = doc.add_table(rows=0, cols=len(table.headings))
            doc_paragraph = doc.add_paragraph()

            paragraph._p.addprevious(doc_table._tbl)
            paragraph._p.addprevious(doc_paragraph._p)

            table.fill_docx(
                doc_table,
                title_style="CAN Table Title",
                heading_style="CAN Table Heading",
                contents_style="CAN Table Contents",
            )

        # TODO: Would rather delete the tag paragraph but that breaks the
        #       template's landscape page format for some reason
        # paragraph._p.getparent().remove(paragraph._p)
        paragraph.clear()
    doc.save(output)
    pass


def _entry_point():
    logging.basicConfig(format="%(asctime)s - %(levelname)s - %(message)s")

    signal.signal(signal.SIGINT, signal.SIG_DFL)

    return main()
